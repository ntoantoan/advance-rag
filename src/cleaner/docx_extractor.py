import os
import re
from xml.etree import ElementTree
import logging
from cleaner.common import Document
from docx import Document as DocxDocument


logger = logging.getLogger(__name__)


class WordExtractor():
    """Load docx files.

    Args:
        file_path: Path to the file to load.
    """

    def __init__(self, file_path: str):
        """Initialize with file path."""
        self.file_path = file_path

        if "~" in self.file_path:
            self.file_path = os.path.expanduser(self.file_path)

    def parse_docx(self, docx_path):
        doc = DocxDocument(docx_path)
        content = []

        hyperlinks_url = None
        url_pattern = re.compile(r"http://[^\s+]+//|https://[^\s+]+")
        for para in doc.paragraphs:
            paragraph_text = ""
            for run in para.runs:
                if run.text and hyperlinks_url:
                    result = f"  [{run.text}]({hyperlinks_url})  "
                    paragraph_text += result
                    hyperlinks_url = None
                else:
                    paragraph_text += run.text
                if "HYPERLINK" in run.element.xml:
                    try:
                        xml = ElementTree.XML(run.element.xml)
                        x_child = [c for c in xml.iter() if c is not None]
                        for x in x_child:
                            if x_child is None:
                                continue
                            if x.tag.endswith("instrText"):
                                if x.text is None:
                                    continue
                                for i in url_pattern.findall(x.text):
                                    hyperlinks_url = str(i)
                    except Exception as e:
                        logger.exception("Failed to parse HYPERLINK xml")
            if paragraph_text.strip():
                content.append(paragraph_text)
        
        return "\n".join(content)

    def __del__(self) -> None:
        if hasattr(self, "temp_file"):
            self.temp_file.close()

    def extract(self) -> list[Document]:
        """Load given path as single page."""
        content = self.parse_docx(self.file_path)
        return [
            Document(
                page_content=content,
                metadata={"source": self.file_path},
            )
        ]